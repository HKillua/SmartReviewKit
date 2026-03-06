"""Word (.docx) document loader with OMML formula extraction."""

from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path
from typing import Any, Dict, List

from src.core.types import Document, DocumentSection
from src.libs.loader.base_loader import BaseLoader
from src.libs.loader.math_utils import postprocess_math

logger = logging.getLogger(__name__)

try:
    from lxml import etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False

_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"

_HEADING_STYLE_RE = re.compile(r"[Hh]eading\s*(\d+)")
_EXERCISE_RE = re.compile(r"例题|练习|习题|Exercise|Problem|思考题", re.IGNORECASE)
_DEFINITION_RE = re.compile(r"定义|Definition", re.IGNORECASE)
_THEOREM_RE = re.compile(r"定理|Theorem|引理|Lemma", re.IGNORECASE)
_EXAMPLE_RE = re.compile(r"[例示]例|Example|示例", re.IGNORECASE)


def _infer_content_type(text: str) -> str:
    preview = text[:300]
    if _EXERCISE_RE.search(preview):
        return "exercise"
    if _DEFINITION_RE.search(preview):
        return "definition"
    if _THEOREM_RE.search(preview):
        return "theorem"
    if _EXAMPLE_RE.search(preview):
        return "example"
    return "concept"


class DocxLoader(BaseLoader):
    """Load .docx files into a Document with OMML formula extraction."""

    def __init__(
        self,
        extract_images: bool = True,
        image_storage_dir: str = "data/images",
    ) -> None:
        self._extract_images = extract_images
        self._image_dir = Path(image_storage_dir)

    def load(self, file_path: str | Path) -> Document:
        path = self._validate_file(file_path)
        if path.suffix.lower() != ".docx":
            raise ValueError(f"File is not a .docx: {path}")

        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError("python-docx is required: pip install python-docx") from exc

        docx_doc = DocxDocument(str(path))
        doc_hash = hashlib.sha256(path.read_bytes()).hexdigest()[:12]

        md_lines: list[str] = []
        images: list[Dict[str, Any]] = []
        image_seq = 0
        sections: list[DocumentSection] = []
        current_section_title = ""
        current_section_level = 1
        current_section_lines: list[str] = []
        current_section_images: list[Dict[str, Any]] = []
        section_start_idx = 0

        def _flush_section(page_idx: int = 0) -> None:
            nonlocal current_section_lines, current_section_images
            if not current_section_lines:
                return
            content = "\n".join(current_section_lines)
            has_formula = "$" in content
            ct = _infer_content_type(current_section_title + " " + content)
            if has_formula and ct == "concept" and len(content) < 400:
                ct = "formula"
            sections.append(DocumentSection(
                title=current_section_title,
                level=current_section_level,
                content=content,
                content_type=ct,
                page_or_slide=page_idx,
                has_formula=has_formula,
                images=list(current_section_images),
            ))
            current_section_lines = []
            current_section_images = []

        for para_idx, para in enumerate(docx_doc.paragraphs):
            heading_level = self._get_heading_level(para)

            if heading_level:
                _flush_section(para_idx)
                current_section_title = para.text.strip()
                current_section_level = heading_level
                prefix = "#" * heading_level
                md_lines.append(f"{prefix} {current_section_title}")
                md_lines.append("")
                section_start_idx = para_idx
                continue

            para_text = self._extract_paragraph_with_math(para)
            para_text = postprocess_math(para_text)

            if self._extract_images and LXML_AVAILABLE:
                for img_data in self._extract_images_from_para(para, docx_doc):
                    image_seq += 1
                    img_id = f"{doc_hash}_{para_idx}_{image_seq}"
                    img_path = self._save_image_bytes(img_data, img_id)
                    placeholder = f"[IMAGE: {img_id}]"
                    para_text += f"\n{placeholder}"
                    img_meta = {"id": img_id, "page": para_idx, "path": img_path}
                    images.append(img_meta)
                    current_section_images.append(img_meta)

            if para_text.strip():
                md_lines.append(para_text)
                current_section_lines.append(para_text)
            md_lines.append("")

        # Tables
        for table_idx, table in enumerate(docx_doc.tables):
            md_table = self._table_to_markdown(table)
            if md_table:
                md_lines.append(md_table)
                md_lines.append("")
                current_section_lines.append(md_table)

        _flush_section()

        title = docx_doc.core_properties.title or path.stem
        full_text = f"# {title}\n\n" + "\n".join(md_lines)

        return Document(
            id=doc_hash,
            text=full_text,
            metadata={
                "source_path": str(path),
                "doc_type": "docx",
                "title": title,
                "paragraph_count": len(docx_doc.paragraphs),
                "images": images,
                "sections": [s.to_dict() for s in sections],
            },
        )

    # ------------------------------------------------------------------
    # Heading detection
    # ------------------------------------------------------------------

    @staticmethod
    def _get_heading_level(para: Any) -> int:
        style_name = para.style.name if para.style else ""
        m = _HEADING_STYLE_RE.match(style_name)
        if m:
            return min(int(m.group(1)), 3)
        return 0

    # ------------------------------------------------------------------
    # Paragraph text + OMML
    # ------------------------------------------------------------------

    def _extract_paragraph_with_math(self, para: Any) -> str:
        if not LXML_AVAILABLE:
            return para.text

        parts: list[str] = []
        for child in para._element:
            tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ""
            ns = etree.QName(child.tag).namespace if isinstance(child.tag, str) else ""

            if ns == _OMML_NS and tag in ("oMath", "oMathPara"):
                from src.libs.loader.math_utils import omml_to_latex
                latex = omml_to_latex(child).strip()
                if latex:
                    if tag == "oMathPara":
                        parts.append(f"$${latex}$$")
                    else:
                        parts.append(f"${latex}$")
            elif ns == _W_NS and tag == "r":
                t_elem = child.find(f"{{{_W_NS}}}t")
                if t_elem is not None and t_elem.text:
                    parts.append(t_elem.text)

        return "".join(parts) if parts else para.text

    # ------------------------------------------------------------------
    # Image extraction
    # ------------------------------------------------------------------

    def _extract_images_from_para(self, para: Any, docx_doc: Any) -> list[bytes]:
        blobs: list[bytes] = []
        if not LXML_AVAILABLE:
            return blobs
        try:
            for drawing in para._element.findall(f".//{{{_WP_NS}}}inline"):
                blip = drawing.find(f".//{{{_A_NS}}}blip")
                if blip is None:
                    continue
                r_embed = blip.get(f"{{{_R_NS}}}embed")
                if r_embed and hasattr(docx_doc.part, "related_parts"):
                    related = docx_doc.part.related_parts.get(r_embed)
                    if related and hasattr(related, "blob"):
                        blobs.append(related.blob)
        except Exception as e:
            logger.debug(f"Image extraction error in paragraph: {e}")
        return blobs

    def _save_image_bytes(self, data: bytes, img_id: str) -> str:
        self._image_dir.mkdir(parents=True, exist_ok=True)
        ext = "png"
        if data[:3] == b"\xff\xd8\xff":
            ext = "jpg"
        elif data[:4] == b"\x89PNG":
            ext = "png"
        out_path = self._image_dir / f"{img_id}.{ext}"
        out_path.write_bytes(data)
        return str(out_path)

    # ------------------------------------------------------------------
    # Table
    # ------------------------------------------------------------------

    @staticmethod
    def _table_to_markdown(table: Any) -> str:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if not rows:
            return ""
        lines = ["| " + " | ".join(rows[0]) + " |"]
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)
